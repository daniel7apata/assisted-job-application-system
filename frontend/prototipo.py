import streamlit as st
from google.cloud import speech
import tempfile
import sounddevice as sd
import numpy as np
import scipy.io.wavfile
import time
import agentes_intencion as ain
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import streamlit as st
import tempfile
import requests
from bs4 import BeautifulSoup
from gtts import gTTS
import time
import json
from urllib.parse import urlparse, parse_qs, urlencode
import uuid
import os
import sys
import base64
import requests
from dotenv import load_dotenv
import noisereduce as nr
import soundfile as sf
import streamlit.components.v1 as components  # <--- 1. IMPORTAR COMPONENTES

load_dotenv()

if getattr(sys, 'frozen', False):
    basedir = sys._MEIPASS
else:
    basedir = os.path.dirname(__file__)

env_path = os.path.join(basedir, '.env')
from dotenv import load_dotenv
load_dotenv(dotenv_path=env_path)

# Corrige la ruta del JSON cuando estás dentro del .exe
gcp_creds = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
if gcp_creds and not os.path.isabs(gcp_creds):
    gcp_creds_path = os.path.join(basedir, gcp_creds)
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = gcp_creds_path
    
    
# URL base de la API
BASE_URL = "https://q10bp8maag.execute-api.us-east-2.amazonaws.com/"

# --- 2. DEFINIR EL CÓDIGO JAVASCRIPT PARA TALKBACK ---
# Este script se inyectará en la página
js_code = """
<script>
// 1. Esperar a que el DOM del iframe esté listo
document.addEventListener('DOMContentLoaded', () => {

    (function() {
        // Esperar a que la API de síntesis de voz esté lista
        if (!'speechSynthesis' in window) {
            console.log('Tu navegador no soporta la síntesis de voz.');
            return;
        }

        // Asegurarse de que la lista de voces se cargue
        window.speechSynthesis.onvoiceschanged = function() {
            // No es crítico, pero es buena práctica
        };

        const synth = window.speechSynthesis;
        let utterance = new SpeechSynthesisUtterance();
        utterance.lang = 'es-ES'; // Español
        utterance.rate = 1.5;     // Velocidad de habla

        // Función principal para hablar
        function speak(text) {
            if (synth.speaking) {
                // Si ya está hablando, cancelar lo anterior para decir lo nuevo
                synth.cancel();
            }
            utterance.text = text;
            synth.speak(utterance);
        }

        // Función para atar los listeners a los inputs
        function attachListeners() {
            // 2. CRÍTICO: Buscar inputs en el documento PADRE
            const inputs = window.parent.document.querySelectorAll('input[type="text"], input[type="password"]');
            
            inputs.forEach(input => {
                // Usamos un data-attribute para no atar el evento múltiples veces
                if (input.dataset.talkbackAttached) {
                    return;
                }
                input.dataset.talkbackAttached = 'true';

                // 1. Al enfocar (focus)
                input.addEventListener('focus', (event) => {
                    // Intentar encontrar la etiqueta (label) del input
                    let labelText = "Cuadro de texto";
                    try {
                        // event.target es el input en el DOM padre
                        const stTextInput = event.target.closest('div[data-testid="stTextInput"], div[data-testid="stPasswordInput"]');
                        if (stTextInput) {
                            const labelElement = stTextInput.querySelector('label');
                            if (labelElement) {
                                labelText = labelElement.textContent;
                            }
                        }
                    } catch (e) {
                        console.error("Error al buscar label:", e);
                    }
                    
                    // Leer el label o un texto genérico
                    speak(labelText);
                });

                // 2. Al escribir (input)
                input.addEventListener('input', (event) => {
                    let char;
                    // event.data tiene el último caracter ingresado
                    if (event.inputType === 'deleteContentBackward') {
                        char = 'borrado';
                    } else {
                        char = event.data;
                    }

                    if (char) {
                        if (event.target.type === 'password') {
                            char = 'asterisco'; // No leer la contraseña
                        }
                        speak(char);
                    }
                });
            });
        }

        // 2. CRÍTICO: Observar el body del documento PADRE
        if (window.parent && window.parent.document.body) {
            const observer = new MutationObserver((mutations) => {
                // Cada vez que algo cambie en la página padre,
                // re-ejecutamos nuestra función para buscar inputs nuevos.
                attachListeners();
            });

            // Observar cambios en todo el body padre
            observer.observe(window.parent.document.body, {
                childList: true, // Observar hijos añadidos o quitados
                subtree: true    // Observar en todo el árbol DOM
            });

            // Ejecución inicial al cargar el script
            attachListeners();
        } else {
            console.error("Talkback no pudo adjuntarse al documento principal.");
        }

    })();

}); // Fin del DOMContentLoaded
</script>
"""
# --- FIN DEL CÓDIGO JAVASCRIPT ---


def speak(mensaje):
    tts = gTTS(text=mensaje, lang='es')
    unique_suffix = f"_{uuid.uuid4().hex}.mp3"
    with tempfile.NamedTemporaryFile(suffix=unique_suffix, delete=False) as tmpfile:
        tts.save(tmpfile.name)
        audio_path = tmpfile.name
    with open(audio_path, "rb") as audio_file:
        st.audio(audio_file.read(), format="audio/mp3", autoplay=True)
        
def pantalla_inicio():
    st.markdown(
        """
        <style>
        .stButton>button {
            position: fixed;
            top: 0;
            left: 0;
            width: 100vw;
            height: 100vh;
            background: #0099ff;
            z-index: 99999;
            display: flex;
            justify-content: center;
            align-items: center;
            font-size: 6vw;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    if st.button("🚀 Iniciar flujo de voz", key="start_fullscreen"):
        st.session_state.inicio_voz = True
        st.rerun()

# --- 3. INYECTAR EL JS EN LA APP ---
# Lo colocamos aquí para que se cargue una vez al inicio
components.html(js_code, height=0)

# Mostrar pantalla de inicio solo si no se ha iniciado el flujo
if 'inicio_voz' not in st.session_state or not st.session_state.inicio_voz:
    pantalla_inicio()
    st.stop()

        
def speech_to_text(phrase_time_limit=55, silence_duration=2.0):
    fs = 16000  # Frecuencia de muestreo
    block_duration = 0.1  # Duración de cada bloque en segundos
    silence_dbfs_threshold = -38  # Umbral en decibeles (ajusta según tu ambiente)
    
    if st.session_state.silence_dbfs_threshold:
        silence_dbfs_threshold = st.session_state.silence_dbfs_threshold
    
    st.info("Grabando... Hable ahora.")
    audio_frames = []
    silence_time = 0
    recording_started = False
    start_time = time.time()
    progress = st.progress(0)

    def rms_to_dbfs(rms):
        # Evita log(0)
        if rms == 0:
            return -100.0
        return 20 * np.log10(rms / 32768.0)

    with sd.InputStream(samplerate=fs, channels=1, dtype='int16') as stream:
        while True:
            block = stream.read(int(fs * block_duration))[0]
            audio_frames.append(block)
            rms = np.sqrt(np.mean(np.square(block.astype(np.float32))))
            dbfs = rms_to_dbfs(rms)
            #st.write(f"Nivel dBFS: {dbfs:.2f}")  # Descomenta para calibrar

            if dbfs > silence_dbfs_threshold:
                recording_started = True
                silence_time = 0
            else:
                if recording_started:
                    silence_time += block_duration

            elapsed = time.time() - start_time
            progress.progress(min(int((elapsed / phrase_time_limit) * 100), 100))
            if elapsed >= phrase_time_limit:
                break
            if recording_started and silence_time >= silence_duration:
                break

    st.info("Grabación finalizada. Procesando audio...")
    recorded_audio = np.concatenate(audio_frames, axis=0)
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmpfile:
        scipy.io.wavfile.write(tmpfile.name, fs, recorded_audio)
        audio_file_path = tmpfile.name

    # --- SUPRESOR DE RUIDO ---
    # Cargar el audio grabado
    audio_data, rate = sf.read(audio_file_path)
    # Suponiendo que los primeros 0.5 segundos son solo ruido
    ruido = audio_data[:int(rate*0.5)]
    # Aplicar reducción de ruido
    audio_limpio = nr.reduce_noise(y=audio_data, sr=rate, y_noise=ruido)
    # Guardar el resultado en un nuevo archivo temporal
    with tempfile.NamedTemporaryFile(suffix="_limpio.wav", delete=False) as tmpfile_limpio:
        sf.write(tmpfile_limpio.name, audio_limpio, rate)
        audio_file_path_limpio = tmpfile_limpio.name
    # Usar el audio limpio para el reconocimiento
    client_speech = speech.SpeechClient()
    with open(audio_file_path_limpio, "rb") as audio_file:
        content = audio_file.read()
    audio_google = speech.RecognitionAudio(content=content)
    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=fs,
        language_code="es-ES",
    )
    try:
        response = client_speech.recognize(config=config, audio=audio_google)
        if not response.results:
            return "No se pudo entender el audio."
        print(response.results[0].alternatives[0].transcript)
        return response.results[0].alternatives[0].transcript
    except Exception as e:
        return f"Error al solicitar resultados: {e}"

# Estados iniciales
if 'page' not in st.session_state:
    st.session_state.page = 'start'
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'user_data' not in st.session_state:
    st.session_state.user_data = {"correo": "", "contrasena": ""}
if 'cv_data' not in st.session_state:
    st.session_state.cv_data = {
        "estudios": "", "skills": "", "experiencia": ""
    }
if 'postulacion' not in st.session_state:
    st.session_state.postulacion = {
        "puesto": "", "modalidad": "", "region": "", "descripcion": "", "postular": None,
        "requiere_cv": None, "cv_adjuntado": False, "datos_confirmados": False,
        "oferta_seleccionada": None, "entrevista_pregunta_idx": 0, "entrevista_respuestas": []
    }
if 'confirmacion' not in st.session_state:
    st.session_state.confirmacion = None


def reset_confirmacion():
    st.session_state.confirmacion = None

def reset_postulacion():
    st.session_state.postulacion = {
        "puesto": "", "modalidad": "", "region": "", "descripcion": "", "postular": None,
        "requiere_cv": None, "cv_adjuntado": False, "datos_confirmados": False,
        "oferta_seleccionada": None, "entrevista_pregunta_idx": 0, "entrevista_respuestas": []
    }
    reset_confirmacion()


def obtener_texto_con_reintento():
    """
    Llama a speech_to_text y, en caso de error, reproduce el audio notificando que no se entendió
    y vuelve a capturar el audio.
    """
    texto = speech_to_text()
    while texto.startswith("Error") or texto.startswith("No se pudo"):
        st.info("No se entendió lo que dijo. Por favor, inténtelo nuevamente.")
        speak("No se entendió lo que dijo. Por favor, inténtelo nuevamente.")
        texto = speech_to_text()
    return texto

# Ejemplo de uso en la página de inicio
def start_page():
    st.title("¿Tiene cuenta creada?")
    
    st.session_state.silence_dbfs_threshold = st.slider(
        "Ajustar Silence dBFS Threshold",
        min_value=-100,
        max_value=0,
        value=-38,
        step=1,
        help="Ajusta el umbral de silencio en dBFS para la función de reconocimiento de voz."
    )
    
    speak("Bienvenido ¿Tiene cuenta creada?")
    
    user_text_voice = obtener_texto_con_reintento()
    assistant_reply = ain.tiene_cuenta(user_text_voice)
    print("STT:", assistant_reply)
    
    if assistant_reply == "R001":
        st.session_state.page = 'login_email'
        reset_confirmacion()
        st.rerun()
    elif assistant_reply == "R002":
        st.session_state.page = 'register_name'
        reset_confirmacion()
        st.rerun()
    else:
        st.rerun()


def listar_usuarios():
    url = BASE_URL + "usuarios"
    response = requests.get(url)
    
    if response.status_code == 200:
        return response.json()
    else:
        return {"error": "No se pudieron obtener los usuarios", "status_code": response.status_code}


def register_name():
    st.title("Crear cuenta - Ingrese su nombre")
    speak("Dime, ¿cuál es tu nombre completo? Presiona Tab para dirigirte al cuadro de texto.")
    nombre = st.text_input("Ingrese su nombre completo:", key="nombre_input")

    if nombre:
        st.session_state.user_data["nombre"] = nombre
        confirm_msg = f"El nombre ingresado es {nombre}, ¿es correcto?"
        st.markdown(f"### {confirm_msg}")
        speak(confirm_msg)
        st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
        respuesta = speech_to_text()
        confirmacion = ain.confirmacion(respuesta)
        st.write(f"Respuesta reconocida: {respuesta}")
        if confirmacion == "R001":
            st.session_state.page = 'register_email'
            st.rerun()
        else:
            st.error("Nombre no confirmado. Por favor, inténtelo nuevamente.")
            st.session_state.user_data["nombre"] = ""
            st.rerun()

def register_email():
    st.title("Crear cuenta - Ingrese su correo electrónico")
    speak("Dime, ¿cuál es tu correo electrónico? Presiona Tab para dirigirte al cuadro de texto.")
    email = st.text_input("Ingrese su correo electrónico:", key="correo_input")
    if email:
        st.session_state.user_data["correo"] = email
        confirm_msg = f"El correo ingresado es {email}, ¿es correcto?"
        st.markdown(f"### {confirm_msg}")
        speak(confirm_msg)
        st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
        respuesta = speech_to_text()
        confirmacion = ain.confirmacion(respuesta)
        st.write(f"Respuesta reconocida: {respuesta}")
        if confirmacion == "R001":
            # Validar si el correo ya existe
            usuarios = listar_usuarios()
            if isinstance(usuarios, list):
                existe = any(u.get("Correo", "").lower() == email.lower() for u in usuarios)
                if existe:
                    mensaje = "El correo ya está registrado. Para recuperar el acceso debe hablar con un administrador."
                    st.error(mensaje)
                    speak(mensaje)
                    st.session_state.user_data["correo"] = ""
                    st.rerun()
                else:
                    st.session_state.page = 'register_password'
                    st.rerun()
            else:
                st.error("No se pudo validar el correo. Intente más tarde.")
                speak("No se pudo validar el correo. Intente más tarde.")
                st.session_state.user_data["correo"] = ""
                st.rerun()
        else:
            st.error("Correo no confirmado. Por favor, inténtelo nuevamente.")
            st.session_state.user_data["correo"] = ""
            

def register_password():
    st.title("Crear cuenta - Ingrese su contraseña")
    speak("Dime, ¿cuál será tu contraseña? Presiona Tab para dirigirte al cuadro de texto.")
    pwd = st.text_input("Ingrese su contraseña:", type="password", key="password_input")
    if pwd:
        st.session_state.user_data["contrasena"] = pwd
        confirm_msg = f"La contraseña ingresada es {pwd}, ¿es correcta?"
        st.markdown(f"### {confirm_msg}")
        speak(confirm_msg)
        st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
        respuesta = speech_to_text()
        confirmacion = ain.confirmacion(respuesta)
        st.write(f"Respuesta reconocida: {respuesta}")
        if confirmacion == "R001":
            st.session_state.logged_in = True
            st.session_state.page = 'register_phone'
            reset_confirmacion()
            st.rerun()
        else:
            st.error("Contraseña no confirmada. Por favor, inténtelo nuevamente.")
            st.session_state.user_data["contrasena"] = ""
            st.rerun()
            
def crear_usuario(nuevo_usuario):
    url = BASE_URL + "usuarios/"
    response = requests.post(url, json=nuevo_usuario)
    return response.json()
            
def register_phone():
    st.title("Crear cuenta - Ingrese su número de teléfono")
    speak("Dime, ¿cuál es tu número de teléfono? Presiona Tab para dirigirte al cuadro de texto.")
    telefono = st.text_input("Ingrese su número de teléfono:", key="telefono_input")
    if telefono:
        st.session_state.user_data["telefono"] = telefono
        confirm_msg = f"El número ingresado es {telefono}, ¿es correcto?"
        st.markdown(f"### {confirm_msg}")
        speak(confirm_msg)
        st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
        respuesta = speech_to_text()
        confirmacion = ain.confirmacion(respuesta)
        st.write(f"Respuesta reconocida: {respuesta}")

        if confirmacion == "R001":
            # Construir el nuevo usuario con los datos recolectados
            nuevo_usuario = {
                "Usuario": st.session_state.user_data["correo"].split("@")[0],
                "Contraseña": st.session_state.user_data["contrasena"],
                "Nombre": st.session_state.user_data.get("nombre", ""),
                "Correo": st.session_state.user_data["correo"],
                "Telefono": st.session_state.user_data["telefono"]
            }
            response = crear_usuario(nuevo_usuario)

            # --- ACCESO COMO DICT (FIX) ---
            payload = None
            if isinstance(response, dict):
                payload = response.get("data", response)
            elif isinstance(response, list) and response:
                payload = response[0]

            user_id = None
            if isinstance(payload, dict):
                # intenta ambas variantes de clave
                user_id = payload.get("ID_Usuario") or payload.get("id_usuario")
                # si la API anida el body como string JSON (AWS API GW típico)
                if not user_id and "body" in payload and isinstance(payload["body"], str):
                    try:
                        body = json.loads(payload["body"])
                        user_id = body.get("ID_Usuario") or body.get("id_usuario")
                    except Exception:
                        pass

            if user_id:
                st.session_state.user_data["id_usuario"] = user_id
                st.success("Cuenta creada exitosamente!")
                speak("Cuenta creada exitosamente!")
                time.sleep(3)
                st.session_state.page = 'post_login_menu'
                st.rerun()
            else:
                st.error("Error al crear el usuario. Respuesta inesperada del servidor.")
                speak("Error al crear el usuario. Por favor, inténtelo nuevamente.")
            # --- FIN FIX ---

        else:
            st.error("Número no confirmado. Por favor, inténtelo nuevamente.")
            st.session_state.user_data["telefono"] = ""
            st.rerun()


def login_email():
    st.title("Iniciar sesión - Ingrese su correo electrónico")
    speak("Ingresa tu correo electrónico. Presiona Tab para dirigirte al cuadro de texto.")
    email = st.text_input("Ingrese su correo electrónico:", key="email_input")
    if email:
        confirm_msg = f"El correo ingresado es {email}, ¿es correcto?"
        st.markdown(f"### {confirm_msg}")
        speak(confirm_msg)
        st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
        
        respuesta = obtener_texto_con_reintento()
        confirmacion = ain.confirmacion(respuesta)
        
        st.write(f"Respuesta reconocida: {respuesta}")
        if confirmacion == "R001":
            st.session_state.user_data["correo"] = email
            st.session_state.page = 'login_password'
            st.rerun()
        else:
            st.error("Correo no confirmado. Por favor, inténtelo nuevamente.")
            st.session_state.user_data["correo"] = ""
            st.rerun()
            
def login(usuario, contraseña):
    url = BASE_URL + "login"
    params = {
        'usuario': usuario,
        'contraseña': contraseña
    }
    response = requests.post(url, params=params)
    return response.json()

def login_password():
    st.title("Iniciar sesión - Ingrese su contraseña")
    speak("Ingresa tu contraseña. Presiona Tab para dirigirte al cuadro de texto.")
    pwd = st.text_input("Ingrese su contraseña:", type="password", key="password_input")
    if pwd:
        confirm_msg = f"La contraseña ingresada es {pwd}, ¿es correcta?"
        st.markdown(f"### {confirm_msg}")
        speak(confirm_msg)
        st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
        
        respuesta = obtener_texto_con_reintento()
        confirmacion = ain.confirmacion(respuesta)
        
        st.write(f"Respuesta reconocida: {respuesta}")
        if confirmacion == "R001":
            st.session_state.user_data["contrasena"] = pwd
            
            response = login(st.session_state.user_data["correo"].split("@")[0], st.session_state.user_data["contrasena"])
            if response and response.get('ID_Usuario'):
                print ("Login response:", response)
                st.session_state.user_data["id_usuario"] = response.get('ID_Usuario')
                st.session_state.logged_in = True
                speak("Inicio de sesión exitoso")
                time.sleep(3)
                st.session_state.page = 'post_login_menu'
                st.rerun()
            else:
                speak("Error al iniciar sesión. Por favor, verifique sus credenciales.")
                st.error("Error al iniciar sesión. Por favor, verifique sus credenciales.")
                time.sleep(4)
                st.session_state.user_data["correo"] = ""
                st.session_state.user_data["contrasena"] = ""
                st.session_state.page = 'login_email'
                st.rerun()

          

        else:
            st.error("Contraseña no confirmada. Por favor, inténtelo nuevamente.")
            st.session_state.user_data["contrasena"] = ""
            st.rerun()
            
def obtener_usuario(usuario_id):
    url = BASE_URL + f"usuarios/{usuario_id}"
    response = requests.get(url)
    return response.json()

def listar_hojas_de_vida():
    url = BASE_URL + "hojas-de-vida"
    response = requests.get(url)
    return response.json()

def post_login_menu():
    datos_usuario = obtener_usuario(st.session_state.user_data["id_usuario"])
    
    st.session_state.user_data["nombre"] = datos_usuario.get('Nombre')
    st.session_state.user_data["correo"] = datos_usuario.get('Correo')
    st.session_state.user_data["telefono"] = datos_usuario.get('Telefono')
    
    st.title("¿Qué desea hacer?")   
    mensaje = "Qué desea hacer? Puede crear hoja de vida, postular a empleo o ver postulaciones"
    speak(mensaje)
    
    user_text_voice = obtener_texto_con_reintento()
    assistant_reply = ain.elegir_modulo(user_text_voice)
    
    if assistant_reply == "R001":
        # Lógica para listar hojas de vida y decidir la página siguiente
        hojas = listar_hojas_de_vida()
        id_usuario = st.session_state.user_data["id_usuario"]
        hojas_usuario = [h for h in hojas if h.get("ID_Usuario") == id_usuario]
        if hojas_usuario:
            # Seleccionar la hoja con el ID_Hojadevida más alto
            hoja_max = max(hojas_usuario, key=lambda h: h.get("ID_Hojadevida", 0))
            st.session_state.cv_data["estudios"] = hoja_max.get("Estudios", "")
            st.session_state.cv_data["skills"] = hoja_max.get("Skills", "")
            st.session_state.cv_data["experiencia"] = hoja_max.get("Experiencia", "")
            st.session_state.page = 'preguntar_cv'
        else:
            st.session_state.page = 'cv_estudios'
            reset_postulacion()
        st.rerun()
    elif assistant_reply == "R002":
        st.session_state.page = 'postulacion_puesto'
        reset_postulacion()
        st.rerun()
    elif assistant_reply == "R003":
        st.session_state.page = 'ver_postulaciones'
        st.rerun()
    elif assistant_reply == "R004":
        st.session_state.page = 'start_page'
        reset_postulacion()
        st.rerun()
    elif assistant_reply == "R000":
        st.rerun()

def preguntar_cv():
    st.title("Desea crear uno desde cero o partir de uno ya creado?")   
    mensaje = "Deseas modificar tu hoja de vida ya creada o crear una nueva desde cero?"
    speak(mensaje)
    user_text_voice = obtener_texto_con_reintento()
    assistant_reply = ain.intencion_nuevo_cv(user_text_voice)
    
    if assistant_reply == "R001":
        st.session_state.page = 'cv_estudios'
        reset_postulacion()
        st.rerun()
    elif assistant_reply == "R002":
        st.session_state.page = 'cv_modificar'
        reset_postulacion()
        st.rerun()
        
def crear_hoja_de_vida(hoja_de_vida_data):
    url = BASE_URL + "hojas_de_vida/"
    response = requests.post(url, json=hoja_de_vida_data)
    return response.json()

def cv_input(label, key, input_type=None):
    st.info(f"Por favor, hable para ingresar {label}...")
    
    mensaje = ""
    
    if key == "estudios":
        mensaje = "Dime, ¿qué estudios tienes?"
    elif key == "skills":
        mensaje = "Dime, ¿qué habilidades técnicas tienes?"
    elif key == "experiencia":
        mensaje = "Dime, ¿qué experiencia laboral tienes?"
        
    speak(mensaje)
    
    value = speech_to_text()
    
    value = ain.intencion_retroceder(value)  
    
    if value == "R001":
        if key == "estudios":
            st.session_state.page = 'post_login_menu'
            reset_postulacion()
        elif key == "skills":
            st.session_state.page = 'cv_estudios'
        elif key == "experiencia":
            st.session_state.page = 'cv_skills'
        st.rerun()
    elif value == "R002":
        st.session_state.page = 'post_login_menu'
        reset_postulacion()
        st.rerun()

    
    if not value.startswith("Error") and not value.startswith("No se pudo"):
        
        assistant_reply = ""
        
        if key == "estudios":
            assistant_reply = ain.mejorar_cv_formacion(value)
        elif key == "skills":
            assistant_reply = ain.mejorar_cv_habilidades(value)
        elif key == "experiencia":
            assistant_reply = ain.mejorar_cv_experiencia(value)
        
        if assistant_reply:
            st.write(f"Reconocido: **{assistant_reply}**")
            
            # Actualizar la información en cv_data
            st.session_state.cv_data[key] = assistant_reply
            
            if key == "experiencia":
                
                datos_cv = {
                    "ID_Usuario": st.session_state.user_data["id_usuario"],
                    "Estudios": st.session_state.cv_data["estudios"],
                    "Skills": st.session_state.cv_data["skills"],
                    "Experiencia": st.session_state.cv_data["experiencia"]
                }
                
                response = crear_hoja_de_vida(datos_cv)
                print("Crear hoja de vida response:", response)
                if response.get('Estudios'):
                    st.success("Hoja de vida creada exitosamente!")
            
            next_pages = {
                "estudios": "cv_skills",
                "skills": "cv_experiencia",
                "experiencia": "cv_resumen",
            }
            st.session_state.page = next_pages[key]
            st.rerun()


def cv_estudios():
    st.title("Crear hoja de vida - Estudios, dime, qué estudios tienes?")
    cv_input("Estudios", "estudios")


def cv_skills():
    st.title("Crear hoja de vida - Habilidades técnicas")      
    cv_input("Skills", "skills")


def cv_experiencia():
    st.title("Crear hoja de vida - Experiencia")     
    cv_input("Experiencia", "experiencia")


def cv_modificar():
    st.title("Modificar hoja de vida existente")
    st.info("Vamos a revisar y modificar tu hoja de vida por secciones. Te leeré cada sección y podrás decir si deseas modificarla o continuar.")
    speak("Vamos a revisar y modificar tu hoja de vida")
    st.session_state.cv_modificar_paso = st.session_state.get("cv_modificar_paso", "estudios")
    if st.session_state.cv_modificar_paso == "estudios":
        cv_modificar_estudios()
    elif st.session_state.cv_modificar_paso == "skills":
        cv_modificar_skills()
    elif st.session_state.cv_modificar_paso == "experiencia":
        cv_modificar_experiencia()
    else:
        st.session_state.page = "cv_resumen"
        st.session_state.cv_modificar_paso = "estudios"
        st.rerun()

def cv_modificar_estudios():
    estudios = st.session_state.cv_data.get("estudios", "")
    mensaje = f"La sección de estudios actual es: {estudios if estudios else 'No hay estudios registrados.'} ¿Deseas modificar esta sección? Di sí para modificar o no para continuar."
    speak(mensaje)
    st.info(mensaje)
    respuesta = obtener_texto_con_reintento()
    confirmacion = ain.conformidad_seccion(respuesta)
    if confirmacion == "R002":
        speak("Dime, ¿cuáles son todos tus estudios actualizados?")
        nuevo = obtener_texto_con_reintento()
        mejorado = ain.mejorar_cv_formacion(nuevo)
        st.session_state.cv_data["estudios"] = mejorado
        st.success("Estudios actualizados.")
        speak("Estudios actualizados.")
        # Volver a leer la sección para confirmar
        cv_modificar_estudios()
    elif confirmacion == "R001":
        st.session_state.cv_modificar_paso = "skills"
        st.rerun()
    else:
        speak("No entendí tu respuesta. Por favor, responde sí para modificar o no para continuar.")
        cv_modificar_estudios()
        
def cv_modificar_skills():
    skills = st.session_state.cv_data.get("skills", "")
    mensaje = f"La sección de habilidades técnicas actual es: {skills if skills else 'No hay habilidades técnicas registradas.'} ¿Deseas modificar esta sección? Di sí para modificar o no para continuar."
    speak(mensaje)
    st.info(mensaje)
    respuesta = obtener_texto_con_reintento()
    confirmacion = ain.conformidad_seccion(respuesta)
    if confirmacion == "R002":
        speak("Dime, ¿cuáles son todas tus habilidades técnicas actualizadas?")
        nuevo = obtener_texto_con_reintento()
        mejorado = ain.mejorar_cv_habilidades(nuevo)
        st.session_state.cv_data["skills"] = mejorado
        st.success("Habilidades técnicas actualizadas.")
        speak("Habilidades técnicas actualizadas.")
        # Volver a leer la sección para confirmar
        cv_modificar_skills()
    elif confirmacion == "R001":
        st.session_state.cv_modificar_paso = "experiencia"
        st.rerun()
    else:
        speak("No entendí tu respuesta. Por favor, responde sí para modificar o no para continuar.")
        cv_modificar_skills()
        
def actualizar_hoja_de_vida(usuario_id, hoja_de_vida_data):
    url = BASE_URL + f"hojas_de_vida/{usuario_id}"
    response = requests.put(url, json=hoja_de_vida_data)
    return response.json()
        
def cv_modificar_experiencia():
    experiencia = st.session_state.cv_data.get("experiencia", "")
    mensaje = f"La sección de experiencia laboral actual es: {experiencia if experiencia else 'No hay experiencia laboral registrada.'} ¿Deseas modificar esta sección? Di sí para modificar o no para continuar."
    speak(mensaje)
    st.info(mensaje)
    respuesta = obtener_texto_con_reintento()
    confirmacion = ain.conformidad_seccion(respuesta)
    if confirmacion == "R002":
        speak("Dime, ¿cuál es toda tu experiencia laboral actualizada?")
        nuevo = obtener_texto_con_reintento()
        mejorado = ain.mejorar_cv_experiencia(nuevo)
        st.session_state.cv_data["experiencia"] = mejorado
        st.success("Experiencia laboral actualizada.")
        speak("Experiencia laboral actualizada.")
        # Volver a leer la sección para confirmar
        cv_modificar_experiencia()
    elif confirmacion == "R001":
        st.session_state.cv_modificar_paso = "fin"
        st.session_state.page = "cv_resumen"
        st.rerun()
    else:
        speak("No entendí tu respuesta. Por favor, responde sí para modificar o no para continuar.")
        cv_modificar_experiencia()

def add_black_line(paragraph):
    # Añade un borde inferior negro al párrafo (línea horizontal)
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Grosor de la línea
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    border.append(bottom)
    pPr.append(border)



def cv_resumen():
    st.title("Resumen y descarga de CV")
    mensaje = "Hoja de vida generada. Haz clic para descargarla... indícame si deseas volver un paso atrás o al inicio"
    speak(mensaje)
    
    # Crear el documento .docx en formato Harvard
    document = Document()
    ptitulo = document.add_paragraph(st.session_state.user_data["nombre"])
    ptitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ptitulo.paragraph_format.line_spacing = Pt(12)
    for run in ptitulo.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    pinfocontacto = document.add_paragraph(f"{st.session_state.user_data["correo"]} - {st.session_state.user_data["telefono"]}")
    pinfocontacto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ptitulo.paragraph_format.line_spacing = Pt(12)
    for run in pinfocontacto.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)
    for section, key in [("Estudios", "estudios"), ("Habilidades técnicas", "skills"), ("Experiencia", "experiencia")]:
        heading = document.add_heading(section, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        add_black_line(heading)
        p = document.add_paragraph(st.session_state.cv_data.get(key, ""))
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Guardar el documento en un buffer en memoria
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    doc_b64 = base64.b64encode(doc_io.read()).decode()

    # CSS para botón grande
    st.markdown(
        """
        <style>
            .fullscreen-btn-descargar-overlay {
                position: fixed;
                top: 0;
                left: 0;
                width: 100vw;
                height: 100vh;
                background: #0099ff;
                color: white;
                z-index: 99999;
                display: flex;
                justify-content: center;
                align-items: center;
            }
            .fullscreen-btn-descargar-overlay a {
                width: 100vw;
                height: 70vh;
                display: flex;
                color: white;
                justify-content: center;
                align-items: center;
                text-decoration: none;
            }
        """,
        unsafe_allow_html=True
    )
    # Botón de descarga personalizado
    st.markdown(
        f"""
        <div class="fullscreen-btn-descargar-overlay">
            <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{doc_b64}" download="CV_Harvard.docx" class="big-download-button">Descargar hoja de vida</a>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    for k, v in st.session_state.cv_data.items():
        st.write(f"**{k.capitalize()}**: {v}")
    
    st.success("Hoja de vida generada. Haz clic para descargarla. En 15 segundo será redirigido al menú principal.")
    time.sleep(15)
    st.session_state.page = 'post_login_menu'
    reset_postulacion()
    st.rerun()


# Pregunta puesto deseado
def postulacion_puesto():
    st.title("Postulación - Puesto deseado")
    
    speak("Dime, cuál es el puesto que deseas postular?")
    puesto = speech_to_text()
    while puesto.startswith("Error") or puesto.startswith("No se pudo"): 
        st.info("No se entendió lo que dijo. Por favor, inténtelo nuevamente.")
        speak("No se entendió lo que dijo. Por favor, inténtelo nuevamente.")
        puesto = speech_to_text()
    
    
    
    #puesto = ain.intencion_retroceder(puesto)  
    #if puesto == "R001":
    #    st.session_state.page = 'post_login_menu'
    #    reset_postulacion()
    #    st.rerun()
    #elif puesto == "R002":
    #    st.session_state.page = 'post_login_menu'
    #    reset_postulacion()
    #    st.rerun()
        
    puesto = ain.identificar_profesion(puesto)    
    st.session_state.postulacion["puesto"] = puesto
    
    # Preparar el mensaje de confirmación y solicitarla por voz
    confirm_msg = f"Puesto de {puesto}, correcto?"
    st.markdown(f"### {confirm_msg}")
    speak(confirm_msg)
    st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
    
    respuesta = speech_to_text()
    respuesta = ain.precisar_puesto(respuesta)
    confirmacion = ain.confirmacion(respuesta)
    
    if confirmacion == "R001":
        st.session_state.page = 'postulacion_modalidad'
        st.rerun()
    else:
        st.info("Puesto no confirmado, por favor intente nuevamente.")
        reset_confirmacion()
        st.rerun()
   

# Pregunta modalidad
def postulacion_modalidad():
    st.title("Postulación - Modalidad de trabajo")
    speak("Dime, qué modalidad de trabajo prefieres? ¿Presencial, virtual o una que permita ambas?")
    
    seleccion = speech_to_text()
    while seleccion.startswith("Error") or seleccion.startswith("No se pudo"): 
        st.info("No se entendió lo que dijo. Por favor, inténtelo nuevamente.")
        speak("No se entendió lo que dijo. Por favor, inténtelo nuevamente.")
        seleccion = speech_to_text()
       
    #seleccion = ain.intencion_retroceder(seleccion)  
        
    #if seleccion == "R001":
    #    st.session_state.page = 'postulacion_puesto'
    #    st.rerun()
    #elif seleccion == "R002":
    #    st.session_state.page = 'post_login_menu'
    #    reset_postulacion()
    #    st.rerun()

    seleccion = ain.identificar_modalidad(seleccion)
            
    st.session_state.postulacion["modalidad"] = seleccion
    
    # Preparar el mensaje de confirmación y solicitar la confirmación por voz
    confirm_msg = f"Deseas modalidad {seleccion} entonces?"
    st.markdown(f"### {confirm_msg}")
    speak(confirm_msg)
    st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
    
    confirmacion_voice = speech_to_text()
    confirmacion_voice = ain.confirmacion(confirmacion_voice)
    
    
    if confirmacion_voice == "R001":
        if seleccion == "R002":
            st.session_state.page = 'postulacion_region'
        else:
            st.session_state.page = 'postulacion_resultados'
        st.rerun()
    else:
        st.info("Modalidad no confirmada, por favor intente nuevamente.")
        reset_confirmacion()
        st.rerun()


# Pregunta región solo si presencial/hibrido
def postulacion_region():
    st.title("Postulación - Región")
    speak("Considerando que la modalidad es presencial/híbrido, ingresa el departamento deseado.")
    
    # Permitir al usuario ingresar la región manualmente
    region = speech_to_text()
    region = ain.intencion_retroceder(region)  
    
    if region == "R001":
        st.session_state.page = 'postulacion_modalidad'
        st.rerun()
    elif region == "R002":
        st.session_state.page = 'post_login_menu'
        reset_postulacion()
        st.rerun()

    
    region = ain.corregir_departamento(region)
    
    if region and region != st.session_state.postulacion["region"]:
        st.session_state.postulacion["region"] = region

    if region:
        # Preparar el mensaje de confirmación y solicitar confirmación por voz
        confirm_msg = f"La región ingresada es {region}, ¿es correcta?"
        st.markdown(f"### {confirm_msg}")
        speak(confirm_msg)
        st.info("Por favor, confirme por voz (diga 'sí' para confirmar o 'no' para reintentar)")
        
        confirmacion_voice = speech_to_text(phrase_time_limit=10, silence_duration=1.0)
        if "sí" in confirmacion_voice.lower() or "si" in confirmacion_voice.lower():
            st.session_state.page = 'postulacion_resultados'
            st.rerun()
        else:
            st.info("Región no confirmada, por favor intente nuevamente.")
            reset_confirmacion()
            st.rerun()

def scrape_computrabajo_filtered():
    # Obtener filtros desde st.session_state.postulacion
    puesto = st.session_state.postulacion["puesto"]
    modalidad = st.session_state.postulacion["modalidad"]
    base = "https://pe.computrabajo.com/"
   
    
    # Transforma el puesto (palabras clave) en formato url (minúsculas, palabras separadas por guiones)
    puesto_str = puesto.strip().lower().replace(" ", "-")
    
    if modalidad == "Virtual":
        # Para modalidad virtual se añade "-en-remoto"
        url_path = f"trabajo-de-{puesto_str}-en-remoto"
        
    else:
        # Para Presencial/hibbrido se usa la región (obligatoria en este caso) y luego "-hibrido-"
        region = st.session_state.postulacion["region"]
        region_str = region.strip().lower().replace(" ", "-")
        url_path = f"trabajo-de-{puesto_str}-en-{region_str}-hibrido"
        
        
    # filtrar solo ofertas para personas con discapacidad (no espefica que tipo, mejor no usar)    
    #url = base + url_path + "?dis=1"
    url = base + url_path
    st.write(f"URL generada: {url}")
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/105.0.0.0 Safari/537.36"
    }
    # Realiza la petición HTTP a la URL indicada con headers
    response = requests.get(url, headers=headers, timeout=30)
    if response.status_code != 200:
        st.error("Error al acceder a la página: " + str(response.status_code))
        return []
    
    soup = BeautifulSoup(response.text, "html.parser")
    links = soup.find_all("a", href=True)
    #ofertas = [link['href'] for link in links if "ofertas-de-trabajo" in link['href'] and link['href'].startswith("/ofertas-de-trabajo/")]
    
    ofertas = ["https://pe.computrabajo.com" + link['href']
           for link in links
           if "ofertas-de-trabajo" in link['href'] and link['href'].startswith("/ofertas-de-trabajo/")]
    
    return ofertas
        
# Mostrar resultados simulados según filtros
def crear_vacante(vacante_data):
    url = BASE_URL + "vacantes/"
    response = requests.post(url, json=vacante_data)
    return response.json()

def extract_oferta_data(link):
    headers = {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/105.0.0.0 Safari/537.36")
    }
    try:
        response = requests.get(link, headers=headers, timeout=30)
        if response.status_code != 200:
            return {}
    except Exception as e:
        st.error(f"Error al acceder a la oferta: {e}")
        return {}
    
    soup = BeautifulSoup(response.text, "html.parser")
    
    titulo = soup.find("h1", class_="fwB fs24 mb5 box_detail w100_m")
    empresa = soup.find("a", class_="dIB fs16 js-o-link")
    condiciones = soup.find("span", class_="tag base mb10")
    descripcion = soup.find("p", class_="mbB")
    requisitos_tags = soup.find_all("li", class_="mb10")
    aptitudes_tags = soup.find_all("span", class_="tag bg_brand_light fc_base mr5 mt10 big")
    beneficios_container = soup.find("div", class_="mt30")
    
    oferta_data = {
        "titulo": titulo.get_text(strip=True) if titulo else "",
        "empresa": empresa.get_text(strip=True) if empresa else "",
        "condiciones": condiciones.get_text(strip=True) if condiciones else "",
        "descripcion": descripcion.get_text(strip=True) if descripcion else "",
        "requisitos": [li.get_text(strip=True) for li in requisitos_tags] if requisitos_tags else [],
        "aptitudes": [span.get_text(strip=True) for span in aptitudes_tags] if aptitudes_tags else [],
        "beneficios": beneficios_container.get_text(strip=True) if beneficios_container else "No se mencionan"
    }
    

    
    return oferta_data


def generar_string_ofertas(ofertas):
    """
    Genera un string con toda la información de todas las ofertas,
    separando cada oferta e indicando los subtítulos pertinentes.
    Si la oferta viene anidada en la clave "data", se utiliza ese diccionario.
    """
    resultado = []
    for idx, oferta in enumerate(ofertas, 1):
        oferta_data = oferta.get("data", oferta)
        partes = [
            f"Oferta {idx}:",
            f"Título: {oferta_data.get('titulo', '')}",
            f"Empresa: {oferta_data.get('empresa', '')}",
            f"Condiciones: {oferta_data.get('condiciones', '')}",
            f"Descripción: {oferta_data.get('descripcion', '')}",
            f"Requisitos: {', '.join(oferta_data.get('requisitos', [])) if oferta_data.get('requisitos') else 'No especificado'}",
            f"Aptitudes: {', '.join(oferta_data.get('aptitudes', [])) if oferta_data.get('aptitudes') else 'No especificado'}",
            f"Beneficios: {oferta_data.get('beneficios', '')}",
            "-" * 40
        ]
        resultado.append('\n'.join(partes))
    return '\n\n'.join(resultado)



def transformar_link(link_original):
    # Extraer el ID de la oferta desde la URL original
    try:
        id_oferta = link_original.split('-')[-1].split('#')[0]
    except IndexError:
        print("❌ No se pudo extraer el ID de la oferta.")
        return None

    # Extraer parámetros de la URL original (como lc)
    parsed_url = urlparse(link_original)
    query_params = parse_qs(parsed_url.fragment)  # en este caso, lc está en el fragmento

    lc = query_params.get('lc', [''])[0]  # obtenemos el valor de lc si existe

    # Construir nuevo enlace con los parámetros fijos + ID
    base_url = "https://candidato.pe.computrabajo.com/match/"
    nuevos_params = {
        'oi': id_oferta,
        'p': '57',
        'idb': '1',
        'd': '32',
    }

    if lc:
        nuevos_params['lc'] = lc

    # Generar URL final
    url_final = base_url + '?' + urlencode(nuevos_params)
    return url_final

def listar_vacantes():
    url = BASE_URL + "vacantes"
    response = requests.get(url)
    return response.json()

def obtener_id_vacante_por_link(link):
    """
    Busca el ID de la vacante cuyo enlace coincide exactamente con el link dado.
    Retorna el ID_Vacantes si lo encuentra, o None si no existe.
    """
    vacantes = listar_vacantes()
    for vacante in vacantes:
        if vacante.get("Enlace") == link:
            return vacante.get("ID_Vacantes")
    return None

def crear_postulacion(postulacion_data):
    url = BASE_URL + "postulaciones/"
    response = requests.post(url, json=postulacion_data)
    return response.json()



def postulacion_resultados():
    st.title("Resultados de búsqueda")
    st.session_state['oferta_index'] = 0
    msg_container = st.empty()
    oferta_container = st.empty()
    speak("Se están buscando las ofertas acorde a sus filtros, por favor, espera...")
    msg_container.write("Se están buscando las ofertas acorde a sus filtros...")

    # Invocar la función que devuelve los links segmentados
    links = scrape_computrabajo_filtered()

    if not links:
        msg_container.warning("No se encontraron ofertas con los criterios ingresados.")
        speak("No se encontraron ofertas con los criterios ingresados. ¿Deseas regresar?")
        alternativa = speech_to_text()
        alternativa = ain.intencion_retroceder(alternativa)
        if alternativa == "R001":
            st.session_state.page = 'postulacion_modalidad'
            st.rerun()
        elif alternativa == "R002":
            st.session_state.page = 'post_login_menu'
            reset_postulacion()
            st.rerun()
        return

    ofertas = []
    for link in links:
        data = extract_oferta_data(link)
        

        oferta_guardar = {
            "Titulo": data.get("titulo", ""),
            "Empresa": data.get("empresa", ""),
            "Condiciones": data.get("condiciones", ""),
            "Descripcion": data.get("descripcion", ""),
            "Region": st.session_state.postulacion["region"],
            "Modalidad": st.session_state.postulacion["modalidad"],
            # Convertir listas a strings separados por coma
            "Requisitos": ", ".join(data.get("requisitos", [])) if isinstance(data.get("requisitos", []), list) else data.get("requisitos", ""),
            "Aptitudes": ", ".join(data.get("aptitudes", [])) if isinstance(data.get("aptitudes", []), list) else data.get("aptitudes", ""),
            "Enlace": link,
            "Beneficios": data.get("beneficios", "")
        }

        
        print(crear_vacante(oferta_guardar))
        
        if data:
            ofertas.append({"data": data, "link": link})

    if not ofertas:
        msg_container.warning("No se encontraron ofertas con datos válidos.")
        speak("No se encontraron ofertas con datos válidos. ¿Deseas regresar?")
        alternativa = speech_to_text()
        alternativa = ain.intencion_retroceder(alternativa)
        if alternativa == "R001":
            st.session_state.page = 'postulacion_modalidad'
            st.rerun()
        elif alternativa == "R002":
            st.session_state.page = 'post_login_menu'
            reset_postulacion()
            st.rerun()
        return

    st.session_state['ofertas'] = ofertas
    st.session_state['oferta_index'] = st.session_state.get('oferta_index', 0)

    if st.session_state['oferta_index'] >= len(ofertas):
        msg_container.warning("Se han revisado todas las ofertas disponibles.")
        speak("Se han revisado todas las ofertas disponibles. ¿Deseas regresar?")
        alternativa = speech_to_text()
        alternativa = ain.intencion_retroceder(alternativa)
        if alternativa == "R001":
            st.session_state.page = 'postulacion_modalidad'
            st.rerun()
        elif alternativa == "R002":
            st.session_state.page = 'post_login_menu'
            reset_postulacion()
            st.rerun()
        return

    # Mostrar la oferta actual
    current_oferta = st.session_state['ofertas'][st.session_state['oferta_index']]

    with oferta_container:
        st.markdown(f"**Oferta {st.session_state['oferta_index'] + 1}:**")
        st.markdown(f"**Título:** {current_oferta['data']['titulo']}")
        st.markdown(f"**Empresa:** {current_oferta['data']['empresa']}")
        st.markdown(f"**Condiciones:** {current_oferta['data']['condiciones']}")
        st.markdown(f"**Descripción:** {current_oferta['data']['descripcion']}")
        st.markdown(
            f"**Requisitos:** {', '.join(current_oferta['data']['requisitos']) if current_oferta['data']['requisitos'] else 'No especificado'}"
        )
        st.markdown(
            f"**Aptitudes:** {', '.join(current_oferta['data']['aptitudes']) if current_oferta['data']['aptitudes'] else 'No especificado'}"
        )
        st.markdown(f"**Beneficios:** {current_oferta['data']['beneficios']}")

    oferta_string = generar_string_ofertas([current_oferta])
    resumen_oferta = ain.resumir_oferta(oferta_string)
    speak(resumen_oferta)

    msg_container.write("¿Te agrada esta oferta? Di 'sí' para seleccionarla o 'no' para pasar a la siguiente.")
    respuesta = speech_to_text()
    confirmacion = ain.elegir_oferta(respuesta)
    if confirmacion == "R001":
        st.session_state.postulacion["oferta_seleccionada"] = st.session_state['oferta_index']
        
        st.session_state.postulacion["id_vacante_elegida"] = obtener_id_vacante_por_link(current_oferta['link'])
        
        postulacion_data = {
            "ID_Usuario": st.session_state.user_data["id_usuario"],
            "ID_Vacantes": st.session_state.postulacion["id_vacante_elegida"]
        }
       
        response_guardar_postulacion = crear_postulacion(postulacion_data)
        print(response_guardar_postulacion)
       
        link_postulacion = transformar_link(current_oferta['link'])
        if link_postulacion:
            st.session_state.link_postulacion = link_postulacion
            msg_container = st.empty()
            oferta_container = st.empty()
            st.session_state.page = 'boton_postular_trabajo'
            st.rerun()
        else:
            msg_container.error("No se pudo generar el enlace directo de postulación.")
            speak("No se pudo generar el enlace directo de postulación.")
    else:
        msg_container.info("Vamos a revisar la siguiente oferta.")
        st.session_state['oferta_index'] += 1
        st.rerun()


def boton_postular_trabajo():
    speak("Para postular, pulse el botón grande que aparece en pantalla. En 5 segundos será redirigido al menú principal.")
    if "link_postulacion" in st.session_state:
        link = st.session_state.link_postulacion
        st.markdown(
            """
            <style>
            .fullscreen-btn-overlay {
                position: fixed;
                top: 0;
                left: 0;
                width: 100vw;
                height: 100vh;
                background: white;
                z-index: 99999;
                display: flex;
                justify-content: center;
                align-items: center;
            }
            .fullscreen-btn-overlay a {
                width: 100vw;
                height: 100vh;
                display: flex;
                justify-content: center;
                align-items: center;
                text-decoration: none;
            }
            .fullscreen-btn-overlay button {
                width: 100vw;
                height: 100vh;
                font-size: 6vw;
                background: #0099ff;
                color: white;
                border: none;
                border-radius: 0;
                cursor: pointer;
                margin: 0;
                padding: 0;
            }
            </style>
            """,
            unsafe_allow_html=True,
        )
        st.markdown(
            f"""
            <div class="fullscreen-btn-overlay">
                <a href="{link}" target="_blank">
                    <button>🚀 Postular a esta oferta</button>
                </a>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.error("Enlace de postulación no disponible.")
    
    time.sleep(15)
    st.session_state.page = 'post_login_menu'
    reset_postulacion()
    st.rerun()

def listar_postulaciones():
    url = BASE_URL + "postulaciones"
    response = requests.get(url)
    return response.json()


def ver_postulaciones():
    st.title("Tus postulaciones")
    usuario_id = st.session_state.user_data["id_usuario"]

    # Obtener todas las postulaciones y vacantes
    postulaciones = listar_postulaciones()
    vacantes = listar_vacantes()

    # Filtrar postulaciones del usuario actual
    postulaciones_usuario = [p for p in postulaciones if p.get("ID_Usuario") == usuario_id]

    if not postulaciones_usuario:
        mensaje = "No tienes postulaciones registradas."
        st.info(mensaje)
        speak(mensaje)
        time.sleep(3)
        st.session_state.page = 'post_login_menu'
        st.rerun()

    resumenes = []
    for postulacion in postulaciones_usuario:
        id_vacante = postulacion.get("ID_Vacantes")
        vacante = next((v for v in vacantes if v.get("ID_Vacantes") == id_vacante), None)
        if vacante:
            titulo = vacante.get("Titulo", "Sin título")
            empresa = vacante.get("Empresa", "Sin empresa")
            descripcion = vacante.get("Descripcion", "Sin descripción")
            st.markdown(f"### {titulo}")
            st.markdown(f"**Empresa:** {empresa}")
            st.markdown(f"**Descripción:** {descripcion}")
            
            info_densa = f"Título: {titulo}. Empresa: {empresa}. Descripción: {descripcion}"
            resumen = ain.resumir_oferta_review(info_densa)
            resumenes.append(resumen)

    # Leer en voz alta el resumen de todas las postulaciones
    if resumenes:
        mensaje = "Estas son tus postulaciones: " + " ".join(resumenes)
        speak(mensaje)
        time.sleep(25)
    else:
        mensaje = "No se encontraron detalles de tus postulaciones."
        st.info(mensaje)
        speak(mensaje)
        time.sleep(5)

    st.session_state.page = 'post_login_menu'
    st.rerun()


# ...existing code...
pages = {
    'start': start_page,
    'register_name': register_name,
    'register_email': register_email,
    'register_phone': register_phone,  # Añade la nueva pantalla aquí
    'register_password': register_password,
    'login_email': login_email,
    'login_password': login_password,
    'post_login_menu': post_login_menu,
    'preguntar_cv': preguntar_cv,
    'cv_modificar': cv_modificar,
    'cv_estudios': cv_estudios,
    'cv_skills': cv_skills,
    'cv_experiencia': cv_experiencia,
    'cv_resumen': cv_resumen,
    'postulacion_puesto': postulacion_puesto,
    'postulacion_modalidad': postulacion_modalidad,
    'postulacion_region': postulacion_region,
    'postulacion_resultados': postulacion_resultados,
    'boton_postular_trabajo': boton_postular_trabajo,
    'ver_postulaciones': ver_postulaciones
}

# Ejecutar página actual
pages.get(st.session_state.page, start_page)()